"""
Build V2 CVs from existing DOCX: keep structure/branding, strengthen AI focus.
Outputs:
  Elios_Scoglio_CV_*_V2.docx (+ PDF if LibreOffice/Word available)
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

CV_DIR = Path(__file__).resolve().parent
VIOLET = RGBColor(0x6D, 0x28, 0xD9)
SLATE = RGBColor(0x33, 0x41, 0x55)
DARK = RGBColor(0x0F, 0x17, 0x2A)

CONFIGS = {
    "Elios_Scoglio_CV_FullStackAI.docx": {
        "output": "Elios_Scoglio_CV_FullStackAI_V2.docx",
        "profile_url": "www.108vision.it/profilo/full-stack-ai",
        "role_line": "Full-Stack Developer · AI / LLM Engineer · Applied AI Systems",
        "summary": (
            "Full-Stack Developer and AI / LLM Engineer with 20+ years of continuous "
            "hands-on coding across C# / .NET, Python, Java, Angular, TypeScript and "
            "mobile stacks — from startup to Italy's #1 ticketing platform "
            "(TicketOne / CTS Eventim Group).\n\n"
            "AI-native builder with two complementary strengths: (1) AI-augmented "
            "software delivery — analysis, architecture, coding, testing, review and "
            "documentation accelerated by reusable prompts, quality gates and human "
            "oversight; (2) applied AI product capabilities — personalised audio/"
            "content generation, retrieval & semantic search, document intelligence "
            "and structured multi-agent workflows designed for real teams and products.\n\n"
            "Ships beyond the prompt layer: structured outputs, context routing, model "
            "selection by task, evaluation, safety boundaries, cost visibility and "
            "fallback behaviour. Rare profile combining deep full-stack ownership with "
            "applied AI engineering used in enterprise and product contexts."
        ),
        "ai_section_title": "AI FOCUS — DEVELOPMENT TOOLING + PRODUCT CAPABILITIES",
        "extra_bullets_after": {
            "AI / LLM Engineering — Production Systems": [
                "Applies AI as a daily engineering accelerator: requirements analysis, architecture options, implementation support, refactoring, test design and technical documentation — always under human review and automated checks.",
                "Designs applied AI capabilities for product and client workflows: generation, retrieval/semantic search, document intelligence and guided agentic flows with validation, retries, fallbacks and cost control.",
            ],
        },
        "skills_ai": (
            "Claude (Opus/Sonnet/Haiku) · AWS Bedrock · Anthropic SDK · Multi-Agent "
            "Orchestration · Prompt Engineering · AI-Augmented SDLC · RAG / Semantic "
            "Search · Document Intelligence · Evaluation & Safety · Python AI"
        ),
        "differentiators": [
            "✦  Dual AI Focus — Builder & Accelerator: Ships applied AI capabilities for products and accelerates software delivery with AI-assisted analysis, coding, testing and review — under human ownership.",
            "✦  True Full-Stack: 20 yrs C#/.NET · 8 yrs Angular/TypeScript · 7 yrs Python · 4 yrs Java Spring Boot · iOS/Android/Ionic — owns the path from data and APIs to UX and AI workflows.",
            "✦  Enterprise-Scale Engineering: Hands-on on Italy's #1 ticketing platform — resilience, distributed tracing and event-driven architecture as daily practice.",
            "✦  SaaS Founder Mindset: Founded and solo-coded a SaaS product to 10,000+ users — understands product, delivery speed and ROI alongside engineering excellence.",
            "✦  Human-Centered Collaborator: NLP Counselor + Yoga Teacher — pairs technical precision with strong communication and psychological safety.",
        ],
    },
    "Elios_Scoglio_CV_SoftwareManager.docx": {
        "output": "Elios_Scoglio_CV_SoftwareManager_V2.docx",
        "profile_url": "www.108vision.it/profilo/software-manager",
        "role_line": "Principal Software Architect · Engineering Manager · Applied AI Governance",
        "summary": (
            "Principal Software Architect and Engineering Manager with 20+ years of "
            "hands-on development and architectural leadership across C# / .NET, Java, "
            "Python, Angular and cloud-native stacks — from startup to enterprise scale.\n\n"
            "Currently heading Software Architecture & Development at Italy's #1 "
            "ticketing platform (TicketOne / CTS Eventim Group): steering microservices "
            "migration, defining DDD bounded contexts, enforcing API-First / "
            "Security-by-Design / Observability standards, and introducing applied AI "
            "into both the engineering system and product-facing capabilities.\n\n"
            "AI focus on two controlled fronts: (1) AI-augmented delivery governance — "
            "where AI accelerates analysis, architecture, coding, testing and review "
            "without replacing quality gates; (2) applied AI product governance — "
            "qualifying use cases, requiring evaluation, safety, cost visibility and "
            "fallback behaviour before scaling. Developer-first architect who never "
            "stopped writing code."
        ),
        "ai_section_title": "AI FOCUS — DELIVERY GOVERNANCE + PRODUCT AI",
        "extra_bullets_after": {
            "Hands-On Development & AI Engineering": [
                "Governs AI-augmented engineering workflows used across analysis, architecture, implementation and QA — with mandatory human review and existing CI controls.",
                "Defines go/no-go criteria for applied AI capabilities in product contexts: evaluation, sensitive-data boundaries, provider dependency handling, cost visibility and operational telemetry.",
            ],
        },
        "skills_ai": (
            "Applied AI Governance · AI-Augmented SDLC · Multi-Agent Workflows · "
            "Prompt Assets & Quality Gates · RAG / Semantic Search · Evaluation & "
            "Safety · Cost Routing · Claude / AWS Bedrock · Python AI"
        ),
        "differentiators": [
            "✦  Developer-First Architect: Writes and reviews production code — not architecture theatre; unblocks teams hands-on.",
            "✦  Dual AI Governance: Accelerates the SDLC with AI while governing applied AI capabilities for users and clients — evaluation, safety, cost and ownership.",
            "✦  Full-Stack Range: 20 yrs C#/.NET · 8 yrs Angular · 7 yrs Python · 4 yrs Java — rare depth across the stack.",
            "✦  Human-Centered Leader: NLP Counselor + Yoga Teacher — builds psychological safety and measurable engagement.",
            "✦  Business & Technology Bridge: Startup CEO + enterprise architect — speaks both ROI and architecture trade-offs fluently.",
        ],
    },
    "Elios_Scoglio_CV_TeamLeader.docx": {
        "output": "Elios_Scoglio_CV_TeamLeader_V2.docx",
        "profile_url": "www.108vision.it/profilo/team-leader",
        "role_line": "Principal Software Engineer · Team Leader · AI-Augmented Delivery",
        "summary": (
            "Hands-on Team Leader and Principal Software Engineer who writes code "
            "alongside the team — not just someone who manages tickets. 20+ years of "
            "continuous development across C# / .NET, Angular, Python, Java, TypeScript "
            "and mobile stacks, from startup to enterprise scale.\n\n"
            "Leads high-performing engineering teams at Italy's #1 ticketing platform "
            "(TicketOne / CTS Eventim Group) while remaining an active contributor. "
            "AI focus: coach the team to use AI as a controlled delivery accelerator "
            "(analysis, implementation, tests, review, documentation) and help ship "
            "applied AI capabilities inside real products — with verification, safety "
            "and ownership.\n\n"
            "Unique trajectory: developer-first career that grew into team leadership "
            "while preserving deep technical depth. Certified NLP Counselor: creates "
            "psychological safety and raises team performance."
        ),
        "ai_section_title": "AI FOCUS — TEAM ENABLEMENT + PRODUCT AI",
        "extra_bullets_after": {
            "Hands-On Development & AI Engineering": [
                "Coaches the team on safe AI-assisted development: bounded context, output verification, tests, peer review and ownership of AI-assisted changes.",
                "Helps deliver applied AI product capabilities (document intelligence, semantic search, generative workflows) as normal software — with evaluation, fallback and clear owners.",
            ],
        },
        "skills_ai": (
            "AI-Augmented Delivery · Team AI Conventions · Prompt Playbooks · "
            "Assisted Review & Test Design · Multi-Agent Workflows · RAG / Semantic "
            "Search · Evaluation & Safety · Claude / AWS Bedrock · Python AI"
        ),
        "differentiators": [
            "✦  Developer-First Team Leader: Writes code every day — leads by example and unblocks the team hands-on.",
            "✦  Dual AI Focus: Raises team capability with AI-assisted delivery while helping ship applied AI features safely — without creating AI dependency.",
            "✦  Human-Centered Leader: NLP Counselor + Yoga Teacher — psychological safety and sustainable performance.",
            "✦  Full-Stack Range: 20 yrs C#/.NET · 8 yrs Angular · 7 yrs Python · 4 yrs Java.",
            "✦  Business Acumen: Startup CEO + enterprise experience — product, ROI and engineering trade-offs.",
        ],
    },
}


def set_run_text(paragraph: Paragraph, text: str) -> None:
    """Replace all paragraph text, clearing split runs/bookmarks reliably."""
    font_name = None
    font_size = None
    bold = None
    if paragraph.runs:
        src = paragraph.runs[0]
        font_name = src.font.name
        font_size = src.font.size
        bold = src.font.bold

    p = paragraph._p
    for child in list(p):
        # Keep only paragraph properties; drop runs, hyperlinks, fields, etc.
        if child.tag != qn("w:pPr"):
            p.remove(child)

    run = paragraph.add_run(text)
    run.font.name = font_name or "Segoe UI"
    run.font.size = font_size or Pt(10)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = SLATE


def replace_containing(paragraph: Paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    # Prefer run-local replace when possible
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    set_run_text(paragraph, full.replace(old, new))
    return True


def find_paragraph(doc: Document, prefix: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def find_exact(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    # python-docx has no insert_after helper: add at end, then move XML after current node
    new_para = paragraph._parent.add_paragraph(text)
    paragraph._element.addnext(new_para._element)
    try:
        new_para.style = paragraph.style
    except Exception:
        pass
    if paragraph.runs:
        src = paragraph.runs[0]
        for run in new_para.runs:
            run.font.name = src.font.name
            run.font.size = src.font.size or Pt(10)
            if src.font.color and src.font.color.rgb:
                run.font.color.rgb = src.font.color.rgb
            else:
                run.font.color.rgb = SLATE
    return new_para


def update_header_role(doc: Document, role_line: str) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            # Role subtitle under the name
            if any(k in text for k in ("Developer", "Architect", "Engineer", "Manager", "Leader", "Builder", "Systems")) and "ELIOS" not in text:
                set_run_text(paragraph, role_line)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)
                    run.font.size = run.font.size or Pt(11)
                return
            if text.startswith("ELIOS SCOGLIO") and ("·" in text or "Engineer" in text):
                set_run_text(paragraph, f"ELIOS SCOGLIO\n{role_line}")
                return


def update_summary(doc: Document, summary: str) -> None:
    title = find_exact(doc, "PROFESSIONAL SUMMARY")
    if title is None:
        raise RuntimeError("PROFESSIONAL SUMMARY heading not found")
    body = None
    found = False
    title_element = title._element
    for paragraph in doc.paragraphs:
        if paragraph._element is title_element:
            found = True
            continue
        if found and paragraph.text.strip():
            body = paragraph
            break
    if body is None:
        raise RuntimeError("PROFESSIONAL SUMMARY body not found")
    set_run_text(body, summary)
    if summary[:60] not in body.text:
        raise RuntimeError(f"Summary replacement failed. Body starts with: {body.text[:160]!r}")


def update_ai_journey_title(doc: Document, title: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("AI & LLM ENGINEERING JOURNEY"):
            set_run_text(paragraph, title)
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK
            return


def update_ai_timeline(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    # Expect one row with 3 cells: Toscano | Aruba | TicketOne
    if not table.rows:
        return
    row = table.rows[0]
    cells = row.cells
    if len(cells) < 3:
        return
    contents = [
        (
            "2019–2024 · Toscano Immobiliare\n"
            "Applied AI in product workflows\n"
            "Document intelligence for property content; predictive valuation models; "
            "generative listing support; embedding-based semantic search — integrated "
            "into multi-tenant product flows."
        ),
        (
            "2024–2025 · Aruba S.p.A.\n"
            "AI-augmented delivery\n"
            "LLM-assisted code quality and change analysis; migration/ops tooling; "
            "AI support inside CI/CD and engineering workflows — always with tests "
            "and human review."
        ),
        (
            "2025–Present · TicketOne\n"
            "Applied AI + SDLC AI\n"
            "Multi-agent workflows (Analyst→Architect→Dev→QA); AI-assisted quality "
            "signals on changes; incident-analysis support; reusable prompt assets "
            "encoding domain and engineering standards."
        ),
    ]
    for cell, text in zip(cells, contents):
        # Clear and rewrite first paragraph; remove extras lightly
        paragraphs = cell.paragraphs
        if not paragraphs:
            cell.text = text
            continue
        set_run_text(paragraphs[0], text)
        for extra in paragraphs[1:]:
            set_run_text(extra, "")
        for run in paragraphs[0].runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = SLATE


def add_extra_bullets(doc: Document, mapping: dict[str, list[str]]) -> None:
    stop_headings = {
        "Architecture & Back-End Development",
        "Architecture & Technical Governance",
        "Team & Process",
        "Team Leadership, Mentoring & Agile Transformation",
        "Team Leadership & Agile Transformation",
        "Crisis Management & Prevention",
        "Crisis Management",
        "Hands-On Development & AI Engineering",
        "AI / LLM Engineering — Production Systems",
        "Aruba S.p.A. (via SCAI Tecno)  |  Jun 2024 – ago 2025",
    }
    for heading, bullets in mapping.items():
        heading_para = find_exact(doc, heading) or find_paragraph(doc, heading)
        if heading_para is None:
            continue
        heading_element = heading_para._element
        last = heading_para
        collecting = False
        for paragraph in doc.paragraphs:
            if paragraph._element is heading_element:
                collecting = True
                continue
            if not collecting:
                continue
            text = paragraph.text.strip()
            if not text:
                continue
            if text in stop_headings and text != heading:
                if last._element is not heading_element:
                    break
            if len(text) > 40:
                last = paragraph
            elif last._element is not heading_element:
                break
        anchor = last
        for bullet in bullets:
            if any(bullet[:60] in p.text for p in doc.paragraphs):
                continue
            anchor = insert_paragraph_after(anchor, bullet)


def update_skills_ai_row(doc: Document, skills_ai: str) -> None:
    if len(doc.tables) < 3:
        return
    table = doc.tables[2]
    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        label = cells[0].text.strip()
        if label.startswith("AI"):
            # rewrite skills cell
            target = cells[1] if len(cells) > 1 else cells[0]
            if target.paragraphs:
                set_run_text(target.paragraphs[0], skills_ai)
                for extra in target.paragraphs[1:]:
                    set_run_text(extra, "")
                for run in target.paragraphs[0].runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = SLATE
            return


def update_differentiators(doc: Document, lines: list[str]) -> None:
    title = find_exact(doc, "UNIQUE DIFFERENTIATORS")
    if title is None:
        return
    diffs: list[Paragraph] = []
    found = False
    title_element = title._element
    for paragraph in doc.paragraphs:
        if paragraph._element is title_element:
            found = True
            continue
        if not found:
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("EU work authorization"):
            break
        if (
            "AI Native" in text
            or "Developer-First" in text
            or "Full-Stack" in text
            or "Human-Centered" in text
            or "Business" in text
            or "SaaS" in text
            or "Enterprise-Scale" in text
            or "Dual AI" in text
            or text.startswith("✦")
        ):
            diffs.append(paragraph)
            continue
        if diffs:
            break
    for idx, data in enumerate(lines):
        if idx < len(diffs):
            set_run_text(diffs[idx], data)
            for run in diffs[idx].runs:
                run.font.size = run.font.size or Pt(10)
                run.font.color.rgb = SLATE
        else:
            insert_paragraph_after(diffs[-1] if diffs else title, data)


def update_contacts(doc: Document, profile_url: str) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "eliosnur@gmail.com" in paragraph.text:
                        set_run_text(
                            paragraph,
                            paragraph.text.replace("eliosnur@gmail.com", "info@108vision.it"),
                        )
                    if "github.com/fuzzy77" in paragraph.text and profile_url not in paragraph.text:
                        set_run_text(
                            paragraph,
                            paragraph.text.replace(
                                "github.com/fuzzy77",
                                f"github.com/fuzzy77\n{profile_url}",
                            ),
                        )


def add_footer_note(doc: Document, profile_url: str) -> None:
    note = doc.add_paragraph()
    run = note.add_run(
        "V2 — AI-focused edition  |  108 Vision  |  www.108vision.it  |  "
        f"info@108vision.it  |  {profile_url}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = VIOLET
    run.italic = True


def convert_to_pdf(docx_path: Path) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert as docx2pdf_convert

        docx2pdf_convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    for binary in ("soffice", "libreoffice"):
        try:
            subprocess.run(
                [binary, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                check=True,
                capture_output=True,
                text=True,
            )
            if pdf_path.exists():
                return pdf_path
        except (FileNotFoundError, subprocess.CalledProcessError):
            continue
    return None


def process(source_name: str, cfg: dict) -> None:
    source = CV_DIR / source_name
    output = CV_DIR / cfg["output"]
    if not source.exists():
        raise FileNotFoundError(source)
    shutil.copy2(source, output)
    doc = Document(output)

    update_header_role(doc, cfg["role_line"])
    update_summary(doc, cfg["summary"])
    update_ai_journey_title(doc, cfg["ai_section_title"])
    update_ai_timeline(doc)
    add_extra_bullets(doc, cfg["extra_bullets_after"])
    update_skills_ai_row(doc, cfg["skills_ai"])
    update_differentiators(doc, cfg["differentiators"])
    update_contacts(doc, cfg["profile_url"])
    add_footer_note(doc, cfg["profile_url"])

    doc.save(output)
    print(f"OK DOCX: {output.name}")

    pdf = convert_to_pdf(output)
    if pdf:
        print(f"OK PDF : {pdf.name}")
    else:
        print(f"WARN   : PDF not generated automatically for {output.name}")


def main() -> int:
    print("Building AI-focused V2 CVs from existing DOCX...\n")
    for source_name, cfg in CONFIGS.items():
        try:
            process(source_name, cfg)
        except Exception as exc:
            print(f"ERROR {source_name}: {exc}", file=sys.stderr)
            return 1
        print()
    print("Done. Files are in tracks/Curriculum/")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
