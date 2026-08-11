"""Patch V2 CVs: info@108vision.it + profile URL, then regenerate PDFs."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

CV_DIR = Path(__file__).resolve().parent
SLATE = RGBColor(0x33, 0x41, 0x55)
VIOLET = RGBColor(0x6D, 0x28, 0xD9)

FILES = {
    "Elios_Scoglio_CV_FullStackAI_V2.docx": "www.108vision.it/profilo/full-stack-ai",
    "Elios_Scoglio_CV_SoftwareManager_V2.docx": "www.108vision.it/profilo/software-manager",
    "Elios_Scoglio_CV_TeamLeader_V2.docx": "www.108vision.it/profilo/team-leader",
}


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    # Prefer run-local
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Rebuild paragraph text
    full = paragraph.text.replace(old, new)
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    run = paragraph.add_run(full)
    run.font.size = Pt(9.5)
    run.font.color.rgb = SLATE
    return True


def patch_doc(path: Path, profile_url: str) -> None:
    doc = Document(path)
    changed = False

    # Header / body text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph, "eliosnur@gmail.com", "info@108vision.it"):
                        changed = True
                    # Ensure profile URL present near github/contacts
                    text = paragraph.text
                    if "github.com/fuzzy77" in text and profile_url not in text:
                        if replace_in_paragraph(paragraph, "github.com/fuzzy77", f"github.com/fuzzy77\n{profile_url}"):
                            changed = True
                        elif profile_url not in paragraph.text:
                            # append as extra paragraph in same cell
                            p = cell.add_paragraph(profile_url)
                            for run in p.runs:
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)
                            changed = True

    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, "eliosnur@gmail.com", "info@108vision.it"):
            changed = True
        if paragraph.text.startswith("V2 — AI-focused edition"):
            desired = (
                f"V2 — AI-focused edition  |  108 Vision  |  www.108vision.it  |  "
                f"info@108vision.it  |  {profile_url}"
            )
            if paragraph.text.strip() != desired:
                p = paragraph._p
                for child in list(p):
                    if child.tag != qn("w:pPr"):
                        p.remove(child)
                run = paragraph.add_run(desired)
                run.font.size = Pt(8)
                run.font.color.rgb = VIOLET
                run.italic = True
                changed = True

    if changed:
        doc.save(path)
        print(f"patched: {path.name}")
    else:
        print(f"no changes: {path.name}")


def export_pdf(docx_path: Path) -> None:
    import win32com.client

    src = str(docx_path.resolve())
    out = str(docx_path.with_suffix(".pdf").resolve())
    tmp = str(docx_path.with_name(docx_path.stem + "_TMP.pdf").resolve())
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(src, ReadOnly=True)
        doc.ExportAsFixedFormat(tmp, 17)
        doc.Close(False)
    finally:
        word.Quit()
    Path(out).unlink(missing_ok=True)
    Path(tmp).replace(out)
    print(f"pdf: {docx_path.with_suffix('.pdf').name}")


def main() -> None:
    for name, profile_url in FILES.items():
        path = CV_DIR / name
        patch_doc(path, profile_url)
        export_pdf(path)


if __name__ == "__main__":
    main()
