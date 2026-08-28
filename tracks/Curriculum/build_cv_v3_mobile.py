"""
Build V3 CVs from V1 DOCX: V2 AI focus + Mobile (React Native/Expo at Aruba+TicketOne,
Blazor Hybrid/MAUI from personal WellBeing project + Toscano, TypeScript, Python).

Mobile attribution (enforced by editorial rules):
  React Native / Expo  →  Aruba S.p.A. + TicketOne
  Blazor Hybrid / MAUI →  Personal WellBeing project + Toscano Immobiliare

Outputs:
  Elios_Scoglio_CV_FullStackAI_V3.docx
  Elios_Scoglio_CV_TeamLeader_V3.docx

Run: python build_cv_v3_mobile.py
Requires: pip install python-docx docx2pdf (docx2pdf optional — for PDF)
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

CV_DIR = Path(__file__).resolve().parent
VIOLET = RGBColor(0x6D, 0x28, 0xD9)
SLATE = RGBColor(0x33, 0x41, 0x55)
DARK = RGBColor(0x0F, 0x17, 0x2A)

# ── V3 configs ────────────────────────────────────────────────────────────────

CONFIGS: dict[str, dict] = {
    "Elios_Scoglio_CV_FullStackAI.docx": {
        "output": "Elios_Scoglio_CV_FullStackAI_V3.docx",
        "profile_url": "www.108vision.it/profilo/full-stack-mobile-ai",
        "role_line": "Full-Stack & Mobile Engineer · React Native / Expo · Blazor Hybrid / MAUI · AI / LLM",
        "summary": (
            "Full-Stack, Mobile and AI/LLM Engineer with 20+ years of continuous "
            "hands-on coding across C# / .NET, Python, Java, Angular, TypeScript and "
            "React Native / Expo — from startup to Italy's #1 ticketing platform "
            "(TicketOne / CTS Eventim Group).\n\n"
            "Mobile across two complementary stacks: React Native / Expo (TypeScript) "
            "for cross-platform iOS / Android delivery at enterprise scale, applied at "
            "Aruba and TicketOne; Blazor Hybrid / .NET MAUI (C#) for Microsoft-stack "
            "native apps — developed independently as a personal wellbeing product and "
            "applied commercially at Toscano Immobiliare. Python end to end — data "
            "preparation, automation scripts, AI tooling and backend services.\n\n"
            "AI-native builder with two complementary strengths: (1) AI-augmented "
            "software delivery — analysis, architecture, coding, testing, review and "
            "documentation accelerated by reusable prompts, quality gates and human "
            "oversight; (2) applied AI product capabilities — personalised generation, "
            "retrieval & semantic search, document intelligence and structured "
            "multi-agent workflows designed for real teams and products.\n\n"
            "Ships beyond the prompt layer: structured outputs, context routing, model "
            "selection by task, evaluation, safety boundaries, cost visibility and "
            "fallback behaviour. Rare profile combining deep full-stack and mobile "
            "ownership with applied AI engineering used in enterprise and product contexts."
        ),
        "ai_section_title": "AI FOCUS — DEVELOPMENT TOOLING + PRODUCT CAPABILITIES",
        # extra_bullets_after keys must match paragraph text that starts with the given prefix.
        # Toscano: finds the company header starting with "Toscano".
        # Aruba: finds the company header starting with "Aruba S.p.A.".
        # TicketOne AI subsection: finds "Hands-On Development & AI Engineering".
        "extra_bullets_after": {
            # TicketOne: AI bullets + React Native mobile (FullStackAI uses this subheading)
            "AI / LLM Engineering — Production Systems": [
                "Applies AI as a daily engineering accelerator: requirements analysis, architecture options, implementation support, refactoring, test design and technical documentation — always under human review and automated checks.",
                "Designs applied AI capabilities for product and client workflows: generation, retrieval/semantic search, document intelligence and guided agentic flows with validation, retries, fallbacks and cost control.",
                "Delivered a React Native Expo (TypeScript) mobile extension for venue operations staff (iOS/Android), integrating with the SPORT backend REST APIs for real-time seat management, event notifications over WebSocket and OAuth authentication via the existing IdentityServer stack.",
                "Applied Python for data pipeline automation, AI prompt orchestration tooling, analytics event processing and operational scripts across the platform.",
            ],
            # Aruba mobile — React Native only
            "Aruba S.p.A.": [
                "Developed a React Native Expo mobile application (TypeScript) for internal cloud operations, enabling field engineers and NOC team to monitor infrastructure status, acknowledge alerts and manage incident response from iOS and Android devices.",
                "Built the EAS Build pipeline integrated into existing GitLab CI/CD; delivered push notifications for critical incidents, offline-first local storage and REST integration against internal platform APIs.",
                "Used Python for backend data-preparation scripts, alert correlation logic and AI-assisted log analysis tooling alongside the main service delivery work.",
            ],
            # Toscano — Blazor Hybrid (personal project applied commercially)
            "Toscano": [
                "Applied Blazor Hybrid / .NET MAUI (C#/Razor) to deliver a native property agent application (iOS + Windows), sharing service layer and domain logic directly with the existing .NET backend — no separate mobile codebase, no context switch for the engineering team.",
                "Brought hands-on Blazor Hybrid expertise acquired through independent product development (personal wellbeing application built and maintained over several years) into a commercial delivery context, reducing ramp-up cost on the stack.",
            ],
        },
        "skills_ai": (
            "Claude (Opus/Sonnet/Haiku) · AWS Bedrock · Anthropic SDK · Multi-Agent "
            "Orchestration · Prompt Engineering · AI-Augmented SDLC · RAG / Semantic "
            "Search · Document Intelligence · Evaluation & Safety · Python AI"
        ),
        "skills_mobile": (
            "React Native · Expo SDK · EAS Build · TypeScript (strict) · "
            "Blazor Hybrid · .NET MAUI · C# / Razor · MVVM · "
            "React Navigation · Zod · Offline-First · Push Notifications · "
            "REST / WebSocket · iOS / Android / Windows"
        ),
        "differentiators": [
            "✦  Builder, not prompter: Ships AI capabilities, mobile apps and enterprise services that teams and products actually use.",
            "✦  Full-Stack + Mobile Range: 20 yrs C#/.NET · 8 yrs Angular/TypeScript · 7 yrs Python · 4 yrs Java · React Native/Expo (enterprise delivery at Aruba + TicketOne) · Blazor Hybrid/MAUI (independent product + Toscano) — two mobile stacks, one engineering standard.",
            "✦  Two-Sided AI Experience: Improves how software is built and designs AI features inside software products — developer accelerator and product AI builder.",
            "✦  Enterprise-Scale Engineering: Hands-on on Italy's #1 ticketing platform — resilience, distributed tracing and event-driven architecture as daily practice.",
            "✦  Human-Centered Collaborator: NLP Counselor + Yoga Teacher — pairs technical precision with strong communication and psychological safety.",
        ],
    },
    "Elios_Scoglio_CV_TeamLeader.docx": {
        "output": "Elios_Scoglio_CV_TeamLeader_V3.docx",
        "profile_url": "www.108vision.it/profilo/team-leader-mobile",
        "role_line": "Technical Team Leader · Principal Software Engineer · React Native · Blazor Hybrid / MAUI · AI",
        "summary": (
            "Hands-on Technical Team Leader and Principal Software Engineer who writes "
            "code alongside the team — not just someone who manages tickets. 20+ years "
            "of continuous development across C# / .NET, Angular, Python, Java, "
            "TypeScript and React Native / Expo mobile stacks, from startup to "
            "enterprise scale.\n\n"
            "Mobile across two stacks: React Native / Expo (TypeScript) for "
            "cross-platform iOS / Android delivery, applied at Aruba and TicketOne; "
            "Blazor Hybrid / .NET MAUI (C#) developed independently as a personal "
            "wellbeing product and applied commercially at Toscano Immobiliare. Coaches "
            "teams on both approaches and makes the trade-off explicit: the right stack "
            "for the context, not the same hammer every time. Python across scripting, "
            "automation and AI tooling.\n\n"
            "Leads high-performing engineering teams at Italy's #1 ticketing platform "
            "(TicketOne / CTS Eventim Group) while remaining an active contributor "
            "across backend, mobile and AI workflows. Coaches engineers on safe "
            "AI-assisted development and helps ship applied AI capabilities inside real "
            "products — with verification, safety and ownership.\n\n"
            "Unique trajectory: developer-first career that grew into team leadership "
            "while preserving deep technical depth. Certified NLP Counselor: creates "
            "psychological safety and raises team performance sustainably."
        ),
        "ai_section_title": "AI FOCUS — TEAM ENABLEMENT + PRODUCT AI",
        "extra_bullets_after": {
            # TicketOne — React Native + Python coaching
            "Hands-On Development & AI Engineering": [
                "Coaches the team on safe AI-assisted development: bounded context, output verification, tests, peer review and ownership of AI-assisted changes.",
                "Helps deliver applied AI product capabilities (document intelligence, semantic search, generative workflows) as normal software — with evaluation, fallback and clear owners.",
                "Provided technical direction for the React Native Expo (TypeScript) mobile extension of the venue operations toolchain (iOS/Android); reviewed code, established TypeScript conventions and integration patterns against SPORT backend REST APIs.",
                "Mentors developers on React Native patterns, Expo SDK lifecycle, offline state management and mobile security — same engineering bar as backend, not a simpler layer.",
                "Applies Python for data pipeline automation, AI prompt orchestration tooling and operational scripts; coaches the team on Python-first tooling as a productivity multiplier.",
            ],
            # Aruba — React Native as hands-on contributor
            "Aruba S.p.A.": [
                "Delivered a React Native Expo mobile application (TypeScript) as a hands-on contributor alongside backend and platform work; established EAS Build in existing GitLab CI/CD.",
                "Coached team members on TypeScript mobile patterns, offline-first architecture and push notification handling; set code review standards for the mobile surface aligned with backend engineering practices.",
                "Introduced Python-based tooling for automation, data preparation and AI-assisted workflows; ensured scripts were tested, documented and team-owned — not black boxes.",
            ],
            # Toscano — Blazor Hybrid (personal project applied commercially)
            "Toscano": [
                "Applied Blazor Hybrid / .NET MAUI (C#/Razor) to deliver a native property agent application (iOS + Windows), sharing service layer and domain logic with the existing .NET backend.",
                "Brought hands-on Blazor Hybrid expertise from independent product development (personal wellbeing application) into a commercial delivery context.",
            ],
        },
        "skills_ai": (
            "AI-Augmented Delivery · Team AI Conventions · Prompt Playbooks · "
            "Assisted Review & Test Design · Multi-Agent Workflows · RAG / Semantic "
            "Search · Evaluation & Safety · Claude / AWS Bedrock · Python AI"
        ),
        "skills_mobile": (
            "React Native · Expo SDK · EAS Build · TypeScript (strict) · "
            "Blazor Hybrid · .NET MAUI · C# / Razor · MVVM · "
            "Offline-First · Push Notifications · REST / WebSocket · "
            "iOS / Android / Windows · Python tooling"
        ),
        "differentiators": [
            "✦  Developer-First Team Leader: Writes code every day across backend, mobile and AI tooling — leads by example and unblocks the team hands-on.",
            "✦  Two Mobile Stacks, One Engineering Standard: React Native/Expo (TypeScript, enterprise delivery) · Blazor Hybrid/MAUI (C#, independent product + Toscano) — coaches the full stack with first-hand experience on both.",
            "✦  Dual AI Focus: Raises team capability with AI-assisted delivery while helping ship applied AI features safely — without creating AI dependency.",
            "✦  Human-Centered Leader: NLP Counselor + Yoga Teacher — psychological safety and sustainable team performance.",
            "✦  Business Acumen: Startup CEO + enterprise experience — product, ROI and engineering trade-offs.",
        ],
    },
}

# ── Helpers ───────────────────────────────────────────────────────────────────


def set_run_text(paragraph: Paragraph, text: str) -> None:
    font_name = None
    font_size = None
    bold = None
    if paragraph.runs:
        src = paragraph.runs[0]
        font_name = src.font.name
        font_size = src.font.size
        bold = src.font.bold

    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)

    run = paragraph.add_run(text)
    run.font.name = font_name or "Segoe UI"
    run.font.size = font_size or Pt(10)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = SLATE


def find_paragraph(doc: Document, prefix: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def find_exact(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_para = paragraph._parent.add_paragraph(text)
    paragraph._element.addnext(new_para._element)
    try:
        new_para.style = paragraph.style
    except Exception:
        pass
    if paragraph.runs:
        src = paragraph.runs[0]
        for run in new_para.runs:
            run.font.name = src.font.name
            run.font.size = src.font.size or Pt(10)
            if src.font.color and src.font.color.rgb:
                run.font.color.rgb = src.font.color.rgb
            else:
                run.font.color.rgb = SLATE
    return new_para


def update_header_role(doc: Document, role_line: str) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if (
                any(
                    k in text
                    for k in (
                        "Developer",
                        "Architect",
                        "Engineer",
                        "Manager",
                        "Leader",
                        "Builder",
                        "Systems",
                    )
                )
                and "ELIOS" not in text
            ):
                set_run_text(paragraph, role_line)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)
                    run.font.size = run.font.size or Pt(11)
                return
            if text.startswith("ELIOS SCOGLIO") and ("·" in text or "Engineer" in text):
                set_run_text(paragraph, f"ELIOS SCOGLIO\n{role_line}")
                return


def update_summary(doc: Document, summary: str) -> None:
    title = find_exact(doc, "PROFESSIONAL SUMMARY")
    if title is None:
        raise RuntimeError("PROFESSIONAL SUMMARY heading not found")
    body = None
    found = False
    title_element = title._element
    for paragraph in doc.paragraphs:
        if paragraph._element is title_element:
            found = True
            continue
        if found and paragraph.text.strip():
            body = paragraph
            break
    if body is None:
        raise RuntimeError("PROFESSIONAL SUMMARY body not found")
    set_run_text(body, summary)


def update_ai_journey_title(doc: Document, title: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("AI & LLM ENGINEERING JOURNEY"):
            set_run_text(paragraph, title)
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK
            return


def update_ai_timeline(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    if not table.rows:
        return
    row = table.rows[0]
    cells = row.cells
    if len(cells) < 3:
        return
    contents = [
        (
            "2019–2024 · Toscano Immobiliare\n"
            "Applied AI + Blazor Hybrid / MAUI\n"
            "Document intelligence for property content; predictive valuation; "
            "generative listing support; semantic search — integrated into multi-tenant "
            "product flows. Blazor Hybrid / .NET MAUI (C#) for native property agent "
            "app (iOS + Windows), sharing domain logic with the .NET backend. Expertise "
            "from personal WellBeing product applied in client delivery."
        ),
        (
            "2024–2025 · Aruba S.p.A.\n"
            "AI-augmented delivery + React Native / Expo\n"
            "React Native Expo (TypeScript) mobile app for cloud ops (iOS + Android); "
            "EAS Build in CI/CD; Python for automation, data-prep and AI-assisted log "
            "analysis tooling — alongside service development and team coordination."
        ),
        (
            "2025–Present · TicketOne\n"
            "Applied AI + SDLC AI + React Native / Expo\n"
            "Multi-agent workflows (Analyst→Architect→Dev→QA); AI quality signals on "
            "changes; React Native Expo (TypeScript) mobile extension for venue "
            "operations staff; Python for data pipelines and AI prompt orchestration."
        ),
    ]
    for cell, text in zip(cells, contents):
        paragraphs = cell.paragraphs
        if not paragraphs:
            cell.text = text
            continue
        set_run_text(paragraphs[0], text)
        for extra in paragraphs[1:]:
            set_run_text(extra, "")
        for run in paragraphs[0].runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = SLATE


def add_extra_bullets(doc: Document, mapping: dict[str, list[str]]) -> None:
    stop_headings = {
        "Architecture & Back-End Development",
        "Architecture & Technical Governance",
        "Team & Process",
        "Team Leadership, Mentoring & Agile Transformation",
        "Team Leadership & Agile Transformation",
        "Crisis Management & Prevention",
        "Crisis Management",
        "Hands-On Development & AI Engineering",
        "AI / LLM Engineering — Production Systems",
        "Aruba S.p.A. (via SCAI Tecno)  |  Jun 2024 – ago 2025",
    }
    for heading, bullets in mapping.items():
        heading_para = find_exact(doc, heading) or find_paragraph(doc, heading)
        if heading_para is None:
            print(f"  WARN: heading not found — '{heading}'")
            continue
        heading_element = heading_para._element
        last = heading_para
        collecting = False
        for paragraph in doc.paragraphs:
            if paragraph._element is heading_element:
                collecting = True
                continue
            if not collecting:
                continue
            text = paragraph.text.strip()
            if not text:
                continue
            if text in stop_headings and text != heading:
                if last._element is not heading_element:
                    break
            if len(text) > 40:
                last = paragraph
            elif last._element is not heading_element:
                break
        anchor = last
        for bullet in bullets:
            if any(bullet[:60] in p.text for p in doc.paragraphs):
                continue
            anchor = insert_paragraph_after(anchor, bullet)


def update_skills_ai_row(doc: Document, skills_ai: str) -> None:
    if len(doc.tables) < 3:
        return
    table = doc.tables[2]
    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        label = cells[0].text.strip()
        if label.startswith("AI"):
            target = cells[1] if len(cells) > 1 else cells[0]
            if target.paragraphs:
                set_run_text(target.paragraphs[0], skills_ai)
                for extra in target.paragraphs[1:]:
                    set_run_text(extra, "")
                for run in target.paragraphs[0].runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = SLATE
            return


def update_skills_mobile_row(doc: Document, skills_mobile: str) -> None:
    """Update existing Mobile row, or append to AI skills row."""
    if len(doc.tables) < 3:
        return
    table = doc.tables[2]
    # Try to find a Mobile row first
    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        label = cells[0].text.strip()
        if label.startswith("Mobile") or "React Native" in label:
            target = cells[1] if len(cells) > 1 else cells[0]
            if target.paragraphs:
                set_run_text(target.paragraphs[0], skills_mobile)
                for run in target.paragraphs[0].runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = SLATE
            return
    # No Mobile row — append mobile note to AI row as a second line
    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        label = cells[0].text.strip()
        if label.startswith("AI"):
            target = cells[1] if len(cells) > 1 else cells[0]
            if target.paragraphs:
                existing = target.paragraphs[0].text
                if "React Native" not in existing:
                    set_run_text(
                        target.paragraphs[0],
                        existing.rstrip() + "\n\nMobile: " + skills_mobile,
                    )
                    for run in target.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = SLATE
            return


def update_differentiators(doc: Document, lines: list[str]) -> None:
    title = find_exact(doc, "UNIQUE DIFFERENTIATORS")
    if title is None:
        return
    diffs: list[Paragraph] = []
    found = False
    title_element = title._element
    for paragraph in doc.paragraphs:
        if paragraph._element is title_element:
            found = True
            continue
        if not found:
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("EU work authorization"):
            break
        if (
            "AI Native" in text
            or "Developer-First" in text
            or "Full-Stack" in text
            or "Human-Centered" in text
            or "Business" in text
            or "SaaS" in text
            or "Enterprise-Scale" in text
            or "Dual AI" in text
            or "Builder" in text
            or "Mobile" in text
            or "Two Mobile" in text
            or text.startswith("✦")
        ):
            diffs.append(paragraph)
            continue
        if diffs:
            break
    for idx, data in enumerate(lines):
        if idx < len(diffs):
            set_run_text(diffs[idx], data)
            for run in diffs[idx].runs:
                run.font.size = run.font.size or Pt(10)
                run.font.color.rgb = SLATE
        else:
            anchor = diffs[-1] if diffs else title
            insert_paragraph_after(anchor, data)


def update_contacts(doc: Document, profile_url: str) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "eliosnur@gmail.com" in paragraph.text:
                        set_run_text(
                            paragraph,
                            paragraph.text.replace("eliosnur@gmail.com", "info@108vision.it"),
                        )
                    if "github.com/fuzzy77" in paragraph.text and profile_url not in paragraph.text:
                        set_run_text(
                            paragraph,
                            paragraph.text.replace(
                                "github.com/fuzzy77",
                                f"github.com/fuzzy77\n{profile_url}",
                            ),
                        )


def add_footer_note(doc: Document, profile_url: str) -> None:
    note = doc.add_paragraph()
    run = note.add_run(
        "V3 — Full-Stack + Mobile (React Native + Blazor Hybrid) + AI  |  108 Vision  |  "
        f"www.108vision.it  |  info@108vision.it  |  {profile_url}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = VIOLET
    run.italic = True


def convert_to_pdf(docx_path: Path) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert as docx2pdf_convert

        docx2pdf_convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    for binary in ("soffice", "libreoffice"):
        try:
            subprocess.run(
                [
                    binary,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(docx_path.parent),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            if pdf_path.exists():
                return pdf_path
        except (FileNotFoundError, subprocess.CalledProcessError):
            continue
    return None


def process(source_name: str, cfg: dict) -> None:
    source = CV_DIR / source_name
    output = CV_DIR / cfg["output"]
    if not source.exists():
        raise FileNotFoundError(f"Source not found: {source}")
    shutil.copy2(source, output)
    doc = Document(output)

    update_header_role(doc, cfg["role_line"])
    update_summary(doc, cfg["summary"])
    update_ai_journey_title(doc, cfg["ai_section_title"])
    update_ai_timeline(doc)
    add_extra_bullets(doc, cfg["extra_bullets_after"])
    update_skills_ai_row(doc, cfg["skills_ai"])
    update_skills_mobile_row(doc, cfg["skills_mobile"])
    update_differentiators(doc, cfg["differentiators"])
    update_contacts(doc, cfg["profile_url"])
    add_footer_note(doc, cfg["profile_url"])

    doc.save(output)
    print(f"  OK DOCX: {output.name}")

    pdf = convert_to_pdf(output)
    if pdf:
        print(f"  OK PDF : {pdf.name}")
    else:
        print(f"  WARN   : PDF not generated for {output.name} (install docx2pdf or LibreOffice)")


def main() -> int:
    print("Building V3 CVs (Full-Stack + Mobile + AI)...\n")
    print("Mobile attribution:")
    print("  React Native / Expo  -> Aruba S.p.A. + TicketOne")
    print("  Blazor Hybrid / MAUI -> Personal WellBeing project + Toscano Immobiliare\n")
    for source_name, cfg in CONFIGS.items():
        print(f"Processing {cfg['output']}...")
        try:
            process(source_name, cfg)
        except Exception as exc:
            print(f"  ERROR: {exc}", file=sys.stderr)
            return 1
        print()
    print("Done. Files are in tracks/Curriculum/")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
