import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CV_DIR = r'c:\Code\Documents\Lavoro\Personale\Vision\tracks\Curriculum'
OUTPUT_DIR = r'c:\Code\Documents\Lavoro\Personale\Vision\aia-website\public\cv'

# Brand colors 108 Vision
VIOLET_700 = RGBColor(0x6D, 0x28, 0xD9)
VIOLET_400 = RGBColor(0xA7, 0x8B, 0xFA)
VIOLET_200 = RGBColor(0xDD, 0xD6, 0xFE)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
SLATE_900 = '0F172A'
SLATE_800 = '1E293B'

OLD_BLUES = ['64B5F6', 'A0C4F8', 'B0C8E8', 'D0E4FF']

cvs = [
    {
        'file': 'Elios_Scoglio_CV_SoftwareManager.docx',
        'profile_url': 'www.108vision.it/profilo/software-manager',
    },
    {
        'file': 'Elios_Scoglio_CV_FullStackAI.docx',
        'profile_url': 'www.108vision.it/profilo/full-stack-ai',
    },
    {
        'file': 'Elios_Scoglio_CV_TeamLeader.docx',
        'profile_url': 'www.108vision.it/profilo/team-leader',
    },
]


def update_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            shd.set(qn('w:fill'), fill_color)


def process_cv(cv_config):
    filepath = os.path.join(CV_DIR, cv_config['file'])
    doc = Document(filepath)

    # === 1. HEADER TABLE (Table 0) ===
    table = doc.tables[0]

    # Update header row backgrounds
    for cell in table.rows[0].cells:
        update_cell_shading(cell, SLATE_900)
    for cell in table.rows[1].cells:
        update_cell_shading(cell, SLATE_800)

    # Update contacts in header (replace email, add website)
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if 'eliosnur@gmail.com' in run.text:
                    run.text = run.text.replace('eliosnur@gmail.com', 'info@108vision.it')
                if 'github.com/fuzzy77' in run.text:
                    run.text = run.text + '\n' + cv_config['profile_url']
                # Recolor contact text to violet-200
                if run.font.color and run.font.color.rgb:
                    if str(run.font.color.rgb) in OLD_BLUES or str(run.font.color.rgb) == 'D0E4FF':
                        run.font.color.rgb = VIOLET_200
                # Increase contact font to 9pt
                if run.font.size and run.font.size.pt < 9:
                    run.font.size = Pt(9.0)

    # Update subtitle color (role line under name)
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    if str(run.font.color.rgb) == 'A0C4F8':
                        run.font.color.rgb = VIOLET_400

    # Update name font size
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if run.font.size and run.font.size.pt == 22.0 and run.font.bold:
                    run.font.size = Pt(24.0)
                    break
            break

    # Update stats colors + font
    for cell in table.rows[1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    hex_c = str(run.font.color.rgb)
                    if hex_c == '64B5F6':
                        run.font.color.rgb = VIOLET_400
                    elif hex_c in ['B0C8E8', 'D0E4FF']:
                        run.font.color.rgb = VIOLET_200
                if run.font.size and run.font.size.pt == 12.0 and run.font.bold:
                    run.font.size = Pt(14.0)
                if run.font.size and run.font.size.pt == 7.5:
                    run.font.size = Pt(8.5)

    # === 2. BODY - increase fonts + recolor ===
    for para in doc.paragraphs:
        for run in para.runs:
            # Recolor old blues
            if run.font.color and run.font.color.rgb:
                hex_c = str(run.font.color.rgb)
                if hex_c in OLD_BLUES:
                    run.font.color.rgb = VIOLET_700

            # Increase font sizes
            if run.font.size:
                current = run.font.size.pt
                if current == 9.0:
                    run.font.size = Pt(10.0)
                elif current == 9.5:
                    run.font.size = Pt(10.5)
                elif current == 8.5:
                    run.font.size = Pt(9.5)
                elif current == 8.0:
                    run.font.size = Pt(9.0)
                elif current == 7.5:
                    run.font.size = Pt(8.5)

    # === 3. OTHER TABLES (AI journey, skills) ===
    for table_idx in range(1, len(doc.tables)):
        t = doc.tables[table_idx]
        for row in t.rows:
            for cell in row.cells:
                # Update dark cell backgrounds
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill.upper() in ['1E2A40', '243350', '1A2540', '2A3A55', '1E3050']:
                            shd.set(qn('w:fill'), SLATE_900)

                for para in cell.paragraphs:
                    for run in para.runs:
                        # Recolor blues to violet
                        if run.font.color and run.font.color.rgb:
                            hex_c = str(run.font.color.rgb)
                            if hex_c in OLD_BLUES:
                                run.font.color.rgb = VIOLET_400
                        # Increase small fonts
                        if run.font.size:
                            current = run.font.size.pt
                            if current <= 8.5:
                                run.font.size = Pt(current + 0.5)

    # === 4. ADD FOOTER BRANDING ===
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.space_before = Pt(12)
    run = footer_para.add_run(
        '108 Vision  |  www.108vision.it  |  info@108vision.it  |  '
        + cv_config['profile_url']
    )
    run.font.size = Pt(8.5)
    run.font.color.rgb = SLATE_500
    run.font.italic = True
    run.font.name = 'Inter'

    # === 5. SAVE ===
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_docx = os.path.join(OUTPUT_DIR, cv_config['file'])
    doc.save(output_docx)
    print(f'  OK: {cv_config["file"]}')


print('Updating CVs with 108 Vision branding...\n')
for cv in cvs:
    try:
        process_cv(cv)
    except Exception as e:
        print(f'  ERROR {cv["file"]}: {e}')

print('\nDOCX files saved to: aia-website/public/cv/')
print('\nNext step: open each .docx in Word/LibreOffice and Export as PDF.')
print('Or run: libreoffice --headless --convert-to pdf *.docx')
