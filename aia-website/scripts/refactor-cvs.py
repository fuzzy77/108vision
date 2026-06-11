"""
Deep refactoring of 3 CVs for 108 Vision branding.
- Rebuilds header with logo + clean layout
- Updates color palette to violet brand
- Increases readability (font sizes, spacing)
- Adds 108 Vision footer branding
- Improves all table styling for modern look
"""
import sys, io, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

CV_DIR = r'c:\Code\Documents\Lavoro\Personale\Vision\tracks\Curriculum'
OUTPUT_DIR = r'c:\Code\Documents\Lavoro\Personale\Vision\aia-website\public\cv'
LOGO_PATH = os.path.join(CV_DIR, 'logo-108-mark.png')

# Brand colors
SLATE_950 = '0F172A'
SLATE_800 = '1E293B'
SLATE_700 = '334155'
SLATE_100 = 'F1F5F9'
VIOLET_700 = '6D28D9'
VIOLET_400 = 'A78BFA'
VIOLET_200 = 'DDD6FE'
VIOLET_50 = 'F5F3FF'
WHITE = 'FFFFFF'

# CV configs
CVS = [
    {
        'file': 'Elios_Scoglio_CV_SoftwareManager.docx',
        'profile_url': 'www.108vision.it/profilo/software-manager',
        'subtitle_override': None,
    },
    {
        'file': 'Elios_Scoglio_CV_FullStackAI.docx',
        'profile_url': 'www.108vision.it/profilo/full-stack-ai',
        'subtitle_override': None,
    },
    {
        'file': 'Elios_Scoglio_CV_TeamLeader.docx',
        'profile_url': 'www.108vision.it/profilo/team-leader',
        'subtitle_override': None,
    },
]


def set_cell_bg(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    """Set cell margins in twips (1/1440 inch)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}'
        f' w:top w:w="{top}" w:type="dxa"'
        f' w:bottom w:w="{bottom}" w:type="dxa"'
        f' w:start w:w="{left}" w:type="dxa"'
        f' w:end w:w="{right}" w:type="dxa"/>'
    )
    # Use OxmlElement approach instead
    tcMar_el = tc.get_or_add_tcPr().find(qn('w:tcMar'))
    if tcMar_el is not None:
        tc.get_or_add_tcPr().remove(tcMar_el)


def remove_table_borders(table):
    """Remove all borders from table for clean look."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def style_header_table(table, cv_config):
    """Completely restyle the header table (Table 0)."""
    remove_table_borders(table)

    # Row 0: Name + contacts header
    row0 = table.rows[0]
    for cell in row0.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        old_shd = tcPr.find(qn('w:shd'))
        if old_shd is not None:
            tcPr.remove(old_shd)
        set_cell_bg(cell, SLATE_950)
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                # Name (large bold white)
                if run.font.bold and run.font.size and run.font.size.pt >= 20:
                    run.font.size = Pt(26)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = 'Segoe UI'
                # Subtitle (violet-400)
                elif run.font.color and run.font.color.rgb and str(run.font.color.rgb) in [VIOLET_400, VIOLET_200, 'A0C4F8', 'DDD6FE']:
                    run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)
                    run.font.size = Pt(11)
                    run.font.name = 'Segoe UI'
                # Contact info
                elif run.font.size and run.font.size.pt <= 10:
                    run.font.color.rgb = RGBColor(0xDD, 0xD6, 0xFE)
                    run.font.size = Pt(9.5)
                    run.font.name = 'Segoe UI'
                    # Replace old email
                    if 'eliosnur@gmail.com' in run.text:
                        run.text = run.text.replace('eliosnur@gmail.com', 'info@108vision.it')
                    # Add profile URL after github
                    if 'github.com' in run.text and 'www.108vision.it' not in run.text:
                        run.text = run.text + '\n' + cv_config['profile_url']

    # Row 1: Stats row
    if len(table.rows) > 1:
        row1 = table.rows[1]
        for cell in row1.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            old_shd = tcPr.find(qn('w:shd'))
            if old_shd is not None:
                tcPr.remove(old_shd)
            set_cell_bg(cell, SLATE_800)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Segoe UI'
                    if run.font.bold:
                        # Stat number
                        run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)
                        run.font.size = Pt(15)
                    else:
                        # Stat label
                        run.font.color.rgb = RGBColor(0xDD, 0xD6, 0xFE)
                        run.font.size = Pt(8.5)


def style_timeline_table(table):
    """Style the AI journey timeline (Table 1)."""
    remove_table_borders(table)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            old_shd = tcPr.find(qn('w:shd'))
            if old_shd is not None:
                tcPr.remove(old_shd)
            set_cell_bg(cell, VIOLET_50)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                for run in para.runs:
                    run.font.name = 'Segoe UI'
                    if run.font.bold:
                        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                        if run.font.size and run.font.size.pt <= 8:
                            run.font.size = Pt(9)
                    else:
                        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                        if run.font.size and run.font.size.pt <= 8:
                            run.font.size = Pt(8.5)


def style_skills_table(table):
    """Style the skills matrix (Table 2)."""
    remove_table_borders(table)

    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            # Clear any existing shading first
            tcPr = cell._tc.get_or_add_tcPr()
            old_shd = tcPr.find(qn('w:shd'))
            if old_shd is not None:
                tcPr.remove(old_shd)

            if ci == 0:
                # Category column: dark bg
                set_cell_bg(cell, SLATE_950)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
                    for run in para.runs:
                        run.font.name = 'Segoe UI'
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9)
                        run.font.bold = True
            else:
                # Skills column: light bg
                set_cell_bg(cell, SLATE_100)
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after = Pt(4)
                    for run in para.runs:
                        run.font.name = 'Segoe UI'
                        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                        run.font.size = Pt(9)


def style_body_paragraphs(doc):
    """Restyle all body paragraphs for readability."""
    for para in doc.paragraphs:
        # Increase spacing
        if para.paragraph_format.space_after is None or para.paragraph_format.space_after < Pt(4):
            para.paragraph_format.space_after = Pt(4)

        for run in para.runs:
            run.font.name = 'Segoe UI'

            # Section headers (bold, larger)
            if run.font.bold and run.font.size and run.font.size.pt >= 11:
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                run.font.size = Pt(12)

            # Subheaders / company names
            elif run.font.bold:
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                if run.font.size:
                    if run.font.size.pt < 10:
                        run.font.size = Pt(10.5)

            # Regular text
            else:
                if run.font.size:
                    current = run.font.size.pt
                    if current <= 9.0:
                        run.font.size = Pt(10)
                    elif current <= 9.5:
                        run.font.size = Pt(10.5)
                    elif current <= 10:
                        run.font.size = Pt(10.5)

                # Recolor old blues to slate
                if run.font.color and run.font.color.rgb:
                    hex_c = str(run.font.color.rgb)
                    if hex_c in ['64B5F6', 'A0C4F8', 'B0C8E8', 'D0E4FF']:
                        run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)


def add_logo_to_header(doc):
    """Add 108 Vision logo to the first cell of the header."""
    if not os.path.exists(LOGO_PATH):
        print('    (logo not found, skipping)')
        return

    table = doc.tables[0]
    # Add logo as first paragraph in row 0 first cell
    cell = table.rows[0].cells[0]
    # Insert new paragraph at beginning
    first_para = cell.paragraphs[0]
    # Add logo before existing content by inserting run with image
    new_para = cell.paragraphs[0]._element
    logo_para = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="100"/></w:pPr></w:p>')
    new_para.addprevious(logo_para)

    # Now add the image via a new paragraph
    from docx.oxml import OxmlElement
    from docx.shared import Inches

    # Actually insert image into first existing paragraph
    run = first_para.runs[0] if first_para.runs else first_para.add_run()
    # Insert image before text - this is complex with python-docx
    # Simpler: add image to a run before the name
    img_para = cell.add_paragraph()
    img_run = img_para.add_run()
    img_run.add_picture(LOGO_PATH, width=Cm(2.5))
    img_para.paragraph_format.space_after = Pt(4)
    img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Move the image paragraph to the top
    cell_element = cell._tc
    img_para_element = img_para._element
    cell_element.remove(img_para_element)
    # Insert before first paragraph
    first_p = cell_element.find(qn('w:p'))
    if first_p is not None:
        first_p.addprevious(img_para_element)


def add_footer_branding(doc, cv_config):
    """Add branded footer."""
    # Add separator
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(16)
    sep_run = sep.add_run('_' * 80)
    sep_run.font.size = Pt(6)
    sep_run.font.color.rgb = RGBColor(0xDD, 0xD6, 0xFE)

    # Footer line
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(8)
    footer.paragraph_format.space_after = Pt(4)

    r1 = footer.add_run('108 Vision')
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    r1.font.name = 'Segoe UI'

    r2 = footer.add_run('  |  www.108vision.it  |  info@108vision.it  |  ')
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r2.font.name = 'Segoe UI'

    r3 = footer.add_run(cv_config['profile_url'])
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    r3.font.name = 'Segoe UI'
    r3.font.italic = True


def process_cv(cv_config):
    """Full refactoring of a CV."""
    filepath = os.path.join(CV_DIR, cv_config['file'])
    doc = Document(filepath)

    print(f'  Processing: {cv_config["file"]}')

    # 1. Style header
    if len(doc.tables) > 0:
        style_header_table(doc.tables[0], cv_config)
        add_logo_to_header(doc)
        print('    - Header styled + logo added')

    # 2. Style timeline
    if len(doc.tables) > 1:
        style_timeline_table(doc.tables[1])
        print('    - Timeline styled')

    # 3. Style skills matrix
    if len(doc.tables) > 2:
        style_skills_table(doc.tables[2])
        print('    - Skills matrix styled')

    # 4. Body paragraphs
    style_body_paragraphs(doc)
    print('    - Body paragraphs restyled')

    # 5. Footer
    add_footer_branding(doc, cv_config)
    print('    - Footer added')

    # 6. Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, cv_config['file'])
    doc.save(output_path)
    print(f'    SAVED: {output_path}')


print('=' * 60)
print('108 Vision CV Deep Refactoring')
print('=' * 60)
print()

for cv in CVS:
    try:
        process_cv(cv)
    except Exception as e:
        import traceback
        print(f'  ERROR: {e}')
        traceback.print_exc()
    print()

print('Done! Open the .docx files in Word and export as PDF.')
